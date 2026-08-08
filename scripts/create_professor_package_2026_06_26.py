"""Create the professor-share package for the current SEBA-XAI paper state.

The package is intentionally share-friendly: Excel, Word, PDF, ZIP, and CSV
artifacts only. It does not upload to Drive by itself; Drive links are recorded
as pending until a connector or browser upload succeeds.
"""

from __future__ import annotations

import csv
import shutil
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "professor_package_2026-06-26"

OVERLEAF_URL = "https://www.overleaf.com/6537138159mjtsgmbgtnqf#204251"
GITHUB_URL = "https://github.com/Rayudualapati-25/SEBA-XAI"
DRIVE_FOLDER_URL = "[Add Google Drive folder link after upload]"
DRIVE_SHEET_URL = "[Add Google Sheet link after upload]"


@dataclass(frozen=True)
class Artifact:
    source: str
    target: str
    section: str
    purpose: str


ARTIFACTS = [
    Artifact(
        "papers/overleaf_ieee_journal/main.pdf",
        "01_paper/SEBA_XAI_Current_Compiled_Paper_2026-06-26.pdf",
        "Paper",
        "Current compiled IEEE-style paper PDF.",
    ),
    Artifact(
        "papers/seba_xai_ieee_journal_overleaf.zip",
        "01_paper/SEBA_XAI_Overleaf_Project_2026-06-26.zip",
        "Paper",
        "Uploadable Overleaf project with updated Introduction and references.",
    ),
    Artifact(
        "papers/reference_contribution_matrix.xlsx",
        "02_research_matrices/Reference_Contribution_Matrix.xlsx",
        "References",
        "Reference matrix explaining what each source contributed or inspired.",
    ),
    Artifact(
        "papers/target_venue_shortlist.csv",
        "02_research_matrices/Target_Venue_Shortlist.csv",
        "Publication planning",
        "Current shortlist of possible venues.",
    ),
    Artifact(
        "papers/final_paper/artifact_to_claim_table.csv",
        "03_claims_and_evidence/Artifact_To_Claim_Table.csv",
        "Claim control",
        "Mapping between paper claims and the artifact evidence supporting them.",
    ),
    Artifact(
        "results/tables/paper_evidence_index.csv",
        "03_claims_and_evidence/Paper_Evidence_Index.csv",
        "Evidence",
        "Index of result tables and evidence used for paper writing.",
    ),
    Artifact(
        "results/tables/paper_table_01_method_comparison.csv",
        "04_result_tables/Paper_Table_01_Method_Comparison.csv",
        "Results",
        "Main method-comparison table for the paper.",
    ),
    Artifact(
        "results/tables/paper_table_02_tamper_detection.csv",
        "04_result_tables/Paper_Table_02_Tamper_Detection.csv",
        "Results",
        "Tamper-detection comparison table.",
    ),
    Artifact(
        "results/tables/paper_table_03_metadata_exposure.csv",
        "04_result_tables/Paper_Table_03_Metadata_Exposure.csv",
        "Results",
        "Metadata-exposure comparison table.",
    ),
    Artifact(
        "results/tables/paper_table_04_latency_storage.csv",
        "04_result_tables/Paper_Table_04_Latency_Storage.csv",
        "Results",
        "Latency and storage overhead table.",
    ),
    Artifact(
        "results/tables/paper_table_05_policy_ablation.csv",
        "04_result_tables/Paper_Table_05_Policy_Ablation.csv",
        "Results",
        "Policy-ablation result table.",
    ),
    Artifact(
        "results/tables/explanation_audit_quality_summary.csv",
        "04_result_tables/Explanation_Audit_Quality_Summary.csv",
        "XAI",
        "XAI trace completeness, counterfactual, stability, and audit-quality summary.",
    ),
    Artifact(
        "results/tables/multi_seed_summary.csv",
        "04_result_tables/Multi_Seed_Summary.csv",
        "Robustness",
        "Multi-seed experiment summary.",
    ),
    Artifact(
        "results/tables/full_grid_per_attack.csv",
        "04_result_tables/Full_Grid_Per_Attack.csv",
        "Attacks",
        "Attack-wise comparison table across defenses.",
    ),
    Artifact(
        "results/tables/nspi_compromised_signer_sensitivity_summary.csv",
        "04_result_tables/NSPI_Compromised_Signer_Sensitivity_Summary.csv",
        "NS-PI",
        "Sensitivity analysis for validly re-signed compromised-signer attack.",
    ),
    Artifact(
        "research_pack/02_literature_matrix.csv",
        "02_research_matrices/Literature_Matrix.csv",
        "Literature",
        "Research literature matrix.",
    ),
    Artifact(
        "research_pack/04_dataset_matrix.csv",
        "02_research_matrices/Dataset_Matrix.csv",
        "Datasets",
        "Dataset discovery matrix.",
    ),
    Artifact(
        "sources/downloaded_research_papers_2026-05-29/download_index.csv",
        "02_research_matrices/Downloaded_Paper_Index.csv",
        "References",
        "Local index of downloaded reference PDFs.",
    ),
]


def clean_output() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    OUT.mkdir(parents=True)
    for folder in [
        "01_paper",
        "02_research_matrices",
        "03_claims_and_evidence",
        "04_result_tables",
        "05_email",
    ]:
        (OUT / folder).mkdir(parents=True, exist_ok=True)


def copy_artifacts() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for artifact in ARTIFACTS:
        source = ROOT / artifact.source
        target = OUT / artifact.target
        target.parent.mkdir(parents=True, exist_ok=True)
        if source.exists():
            shutil.copy2(source, target)
            status = "Included"
            size = str(target.stat().st_size)
        else:
            status = "Missing locally"
            size = ""
        rows.append(
            {
                "section": artifact.section,
                "file_name": Path(artifact.target).name,
                "local_package_path": str(target),
                "source_path": str(source),
                "purpose": artifact.purpose,
                "status": status,
                "size_bytes": size,
                "drive_link": "[Add file link after upload]",
            }
        )
    return rows


def read_csv_preview(path: Path, max_rows: int = 8) -> list[list[str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as handle:
        reader = csv.reader(handle)
        return [row for _, row in zip(range(max_rows), reader)]


def add_sheet(wb: Workbook, title: str, headers: list[str], rows: list[list[str]]) -> None:
    ws = wb.create_sheet(title=title)
    ws.append(headers)
    for row in rows:
        ws.append(row)

    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(wrap_text=True, vertical="top")

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    widths = {
        1: 24,
        2: 42,
        3: 58,
        4: 58,
        5: 70,
        6: 22,
        7: 18,
        8: 46,
    }
    for idx in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(idx)].width = widths.get(idx, 28)
    ws.freeze_panes = "A2"


def create_workbook(artifact_rows: list[dict[str, str]]) -> Path:
    wb = Workbook()
    wb.remove(wb.active)

    add_sheet(
        wb,
        "01_Links",
        ["Item", "Link", "Status", "Notes"],
        [
            ["Overleaf project", OVERLEAF_URL, "Available", "Professor asked to use IEEE journal template and write there."],
            ["Google Drive folder", DRIVE_FOLDER_URL, "Pending upload", "Fill after Drive upload succeeds."],
            ["Google Sheet package index", DRIVE_SHEET_URL, "Pending upload", "Fill after native Google Sheet import succeeds."],
            ["GitHub repository", GITHUB_URL, "Available", "Useful for code and reproducibility. Local branch may still need push credentials."],
            ["Local package folder", str(OUT), "Created locally", "Upload this folder contents to Drive."],
        ],
    )

    add_sheet(
        wb,
        "02_Work_Done",
        ["Area", "What was completed", "Evidence/file", "Boundary"],
        [
            [
                "Research framing",
                "Narrowed the work to secure, explainable, blockchain-audited access governance for inter-agency police data sharing.",
                "README.md; CONTRIBUTION.md; papers/final_paper/claim_control_memo.md",
                "Not a CCTNS/ICJS replacement and not a crime-prediction paper.",
            ],
            [
                "Professor comments",
                "Updated Introduction to add non-India-only framing, recent research, PSNI example, realistic roles, and challenge-to-contribution mapping.",
                "papers/overleaf_ieee_journal/sections/introduction.tex; reports/iteration/iter_054_supervisor_comment_polish.md",
                "Needs supervisor review before freezing wording.",
            ],
            [
                "Prototype",
                "Implemented synthetic access requests, RBAC baseline, ABAC/PBAC policy oracle, XAI artifacts, audit logs, blockchain-style audit, and off-chain commitments.",
                "src/seba/; prototype/synthetic_access_sim/; results/tables/",
                "Synthetic workload only, not real police records.",
            ],
            [
                "Experiments",
                "Ran baselines, ablations, ordinary tamper tests, compromised-signer tests, NS-PI drift tests, workload stress tests, latency/storage, and XAI audit-quality checks.",
                "results/tables/*.csv; results/FINDINGS.md",
                "No SOTA or deployment claim.",
            ],
            [
                "Paper draft",
                "Created IEEE-style Overleaf source with Introduction, Related Work, Methodology, Results/Discussion, Limitations, and Conclusion.",
                "papers/overleaf_ieee_journal/main.tex; papers/overleaf_ieee_journal/main.pdf",
                "Requires final writing polish, supervisor feedback, and formatting cleanup.",
            ],
        ],
    )

    add_sheet(
        wb,
        "03_Professor_Comments",
        ["Comment/request", "Action taken", "Where changed", "Status"],
        [
            ["Change ICJS sentence", "Reworded ICJS as data exchange between CCTNS and criminal-justice pillars.", "introduction.tex", "Done"],
            ["Do not confine to India perspective only", "Added international framing and recent works.", "introduction.tex", "Done"],
            ["Mention recent 2-3 current research works", "Added recent access-control, ABAC, evidence-protection, criminal-justice data-protection, and XAI citations.", "introduction.tex; references.bib", "Done"],
            ["Include example from reputed source", "Added PSNI breach example with Guardian and ICO references.", "introduction.tex; references.bib", "Done"],
            ["Explain realistic roles", "Added investigating officer, ranked/supervisory officer, forensic expert, lab specialist, prosecutor/court authority, and auditor.", "introduction.tex", "Done"],
            ["Each challenge should have one contribution", "Split Practical Challenges and Research Contributions and stated the one-to-one ordering.", "introduction.tex", "Done"],
        ],
    )

    add_sheet(
        wb,
        "04_Files_In_Package",
        ["Section", "File name", "Local package path", "Source path", "Purpose", "Status", "Size bytes", "Drive link"],
        [[row[key] for key in ["section", "file_name", "local_package_path", "source_path", "purpose", "status", "size_bytes", "drive_link"]] for row in artifact_rows],
    )

    add_sheet(
        wb,
        "05_Evidence_Summary",
        ["Evidence area", "What it supports", "Important limitation"],
        [
            ["Method comparison", "Shows baseline/proposed behavior on the synthetic access-governance benchmark.", "Does not prove real police deployment performance."],
            ["Tamper detection", "Shows which audit methods detect ordinary tamper cases.", "Integrity-only methods can miss validly re-signed corruption."],
            ["Metadata exposure", "Shows what audit designs expose in metadata.", "Does not implement full cryptographic metadata privacy."],
            ["Latency/storage", "Gives prototype overhead measurements.", "Local synthetic benchmark, not production infrastructure."],
            ["Policy ablation", "Shows effect of removing policy components.", "Depends on declared synthetic policy oracle."],
            ["XAI audit quality", "Measures trace completeness, counterfactual coverage/validity, explanation stability, decisive-attribute coverage, and audit reconstruction.", "Structured traces are stronger than natural-language explanation text."],
        ],
    )

    add_sheet(
        wb,
        "06_Boundaries",
        ["Do not claim", "Correct wording"],
        [
            ["Real CCTNS/ICJS deployment", "CCTNS/ICJS-compatible synthetic research prototype."],
            ["Real FIR or police-record testing", "Synthetic access-control workload inspired by inter-agency data sharing."],
            ["Legal compliance", "Discusses legal/ethical risks; does not prove compliance."],
            ["Production security", "Evaluates selected threat cases in a prototype."],
            ["State-of-the-art performance", "Compares baselines and ablations inside this benchmark."],
            ["Crime/suspect prediction", "Explains and audits access decisions, not criminal behavior."],
        ],
    )

    add_sheet(
        wb,
        "07_Next_Actions",
        ["Priority", "Action", "Why it matters", "Owner/status"],
        [
            ["1", "Get supervisor feedback on problem statement and challenge-contribution mapping.", "This decides final introduction structure.", "Student/professor"],
            ["2", "Upload package folder and package index sheet to Google Drive.", "Professor requested Drive links rather than attachments.", "Pending connector/browser upload"],
            ["3", "Sync refreshed Overleaf zip/source to Overleaf.", "Ensures live Overleaf has the latest Introduction fixes.", "Pending live Overleaf check"],
            ["4", "Clean layout warnings before submission.", "Improves IEEE formatting quality.", "Pending"],
            ["5", "Run final reproduction before freezing results.", "Prevents paper claims from drifting away from evidence.", "Pending final submission sprint"],
        ],
    )

    email_body = build_email_body()
    add_sheet(
        wb,
        "08_Email_Draft",
        ["Field", "Content"],
        [["Subject", "Updated research package and paper draft links for review"], ["Body", email_body]],
    )

    add_sheet(
        wb,
        "09_Result_Table_Preview",
        ["Source file", "Preview rows"],
        [
            ["paper_table_01_method_comparison.csv", "\n".join(", ".join(row) for row in read_csv_preview(ROOT / "results/tables/paper_table_01_method_comparison.csv"))],
            ["paper_table_02_tamper_detection.csv", "\n".join(", ".join(row) for row in read_csv_preview(ROOT / "results/tables/paper_table_02_tamper_detection.csv"))],
            ["explanation_audit_quality_summary.csv", "\n".join(", ".join(row) for row in read_csv_preview(ROOT / "results/tables/explanation_audit_quality_summary.csv"))],
        ],
    )

    path = OUT / "SEBA_XAI_Professor_Package_Index_2026-06-26.xlsx"
    wb.save(path)
    return path


def build_email_body() -> str:
    return f"""Dear Sir,

I have updated the current research package and the paper draft according to your latest comments.

I am mentioning the required links below:

Overleaf project:
{OVERLEAF_URL}

Google Drive folder:
{DRIVE_FOLDER_URL}

Google Sheet package index:
{DRIVE_SHEET_URL}

GitHub repository:
{GITHUB_URL}

Work completed in the latest update:
1. I changed the ICJS sentence and made the wording more accurate.
2. I revised the Introduction so the paper is not limited only to the India perspective.
3. I added recent related work on blockchain-based access control, privacy-preserving ABAC, blockchain evidence protection, criminal-justice data protection, and XAI in law enforcement.
4. I added a real police-data governance example using the PSNI breach, with reputed media and official ICO references.
5. I added realistic roles such as investigating officer, ranked/supervisory police officer, forensic expert, laboratory specialist, prosecutor/court authority, and auditor.
6. I separated the practical challenges and research contributions, and mapped each challenge to one contribution.
7. I refreshed the Overleaf project zip and verified that the LaTeX source compiles locally.

Current research package contents:
1. Current compiled paper PDF.
2. Updated Overleaf project zip.
3. Package index Excel sheet with links, work done, professor comments, evidence files, boundaries, and next actions.
4. Reference contribution matrix.
5. Claim-to-evidence table.
6. Main result tables for method comparison, tamper detection, metadata exposure, latency/storage, policy ablation, and XAI audit quality.

Important scope I am keeping clear:
1. This is not a replacement for CCTNS or ICJS.
2. The current experiments use synthetic access requests, not real police records.
3. The paper does not claim legal compliance, production deployment, or state-of-the-art crime prediction.
4. The paper focuses on secure, explainable, blockchain-audited access governance for sensitive inter-agency record requests.

I request your feedback mainly on:
1. Whether the problem statement is now clear enough.
2. Whether the challenge-to-contribution mapping is acceptable.
3. Whether the Introduction flow is suitable for the paper.
4. Whether the current contribution points need to be narrowed or strengthened before I continue polishing the paper.

Thanks and Regards,
Alapati Venkat Rayudu
"""


def create_email_docx() -> Path:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    doc.add_heading("Professor Email Draft", level=1)
    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    doc.add_paragraph("Subject: Updated research package and paper draft links for review")
    doc.add_paragraph("")
    for block in build_email_body().split("\n\n"):
        paragraph = doc.add_paragraph()
        paragraph.add_run(block)
    path = OUT / "05_email" / "Professor_Email_Draft_2026-06-26.docx"
    doc.save(path)
    return path


def main() -> None:
    clean_output()
    rows = copy_artifacts()
    workbook = create_workbook(rows)
    email_doc = create_email_docx()
    print(f"Created package: {OUT}")
    print(f"Created workbook: {workbook}")
    print(f"Created email draft: {email_doc}")
    missing = [row for row in rows if row["status"] != "Included"]
    if missing:
        print("Missing artifacts:")
        for row in missing:
            print(f"- {row['source_path']}")


if __name__ == "__main__":
    main()
