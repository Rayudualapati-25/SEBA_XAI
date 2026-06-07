"""Create a professor-share folder with only Word and Excel files.

The generated package intentionally excludes source code, implementation logs,
test outputs, and prototype learning material. It keeps research framing,
literature, datasets, gap, architecture, methodology, ethics, paper outline,
and introduction-planning material.
"""

from __future__ import annotations

import csv
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "professor_share_research_only_2026-05-29"

DOC_SETS = [
    (
        "01_Research_Overview_Problem_Gap.docx",
        "Research Overview, Problem Understanding, and Gap",
        [
            ROOT / "research_brief.md",
            ROOT / "00_problem_understanding.md",
            ROOT / "05_research_gap.md",
        ],
    ),
    (
        "02_Literature_Review_Existing_Models.docx",
        "Literature Review and Existing Related Systems",
        [
            ROOT / "01_literature_review.md",
            ROOT / "14_existing_models_for_testing.md",
        ],
    ),
    (
        "03_Dataset_Study.docx",
        "Dataset Study and Publication Suitability",
        [
            ROOT / "03_datasets.md",
        ],
    ),
    (
        "04_Architecture_Methodology_Evaluation.docx",
        "Architecture, Methodology, and Evaluation Plan",
        [
            ROOT / "06_proposed_architecture.md",
            ROOT / "07_methodology.md",
            ROOT / "08_experiment_plan.md",
            ROOT / "09_evaluation_metrics.md",
        ],
    ),
    (
        "05_Ethics_Paper_Outline_Learning_Plan.docx",
        "Ethics, Paper Outline, and Learning Plan",
        [
            ROOT / "10_ethics_security_legal.md",
            ROOT / "11_paper_outline.md",
            ROOT / "12_five_day_learning_plan.md",
        ],
    ),
    (
        "06_Introduction_Writing_Package.docx",
        "Introduction Writing Package",
        [
            ROOT / "papers/final_paper/introduction/introduction_control_document.md",
            ROOT / "papers/final_paper/introduction/introduction_skeleton.md",
            ROOT / "papers/final_paper/introduction/introduction_draft_v0.md",
            ROOT / "papers/final_paper/introduction/reviewer_objection_checklist.md",
            ROOT / "papers/final_paper/introduction/20_day_introduction_plan.md",
        ],
    ),
    (
        "07_Weekly_Update_To_Professor.docx",
        "Weekly Update to Professor",
        [
            ROOT / "professor_update_2026-05-27.md",
        ],
    ),
]


def clean_text(text: str) -> str:
    """Keep the share package research-only and remove implementation-specific wording."""
    text = text.replace("\r\n", "\n")
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    replacements = {
        "prototype": "research model",
        "Prototype": "Research model",
        "PROTOTYPE": "RESEARCH MODEL",
        "codebase": "research workspace",
        "Codebase": "Research workspace",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    blocked_fragments = [
        "src/seba",
        "prototype/",
        "synthetic_access_sim",
        "pytest",
        "make test",
        "make reproduce",
        "Dockerfile",
        "REPRODUCE.md",
        "SESSION_HANDOFF.md",
    ]
    kept: list[str] = []
    for line in text.splitlines():
        if any(fragment in line for fragment in blocked_fragments):
            continue
        kept.append(line)
    return "\n".join(kept)


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_title(doc: Document, title: str) -> None:
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Prepared for faculty review | Research-only package | 29 May 2026")
    run.italic = True


def is_table_separator(line: str) -> bool:
    parts = [p.strip() for p in line.strip().strip("|").split("|")]
    return bool(parts) and all(re.fullmatch(r":?-{3,}:?", p or "---") for p in parts)


def collect_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i])
        i += 1

    rows: list[list[str]] = []
    for idx, line in enumerate(table_lines):
        if idx == 1 and is_table_separator(line):
            continue
        cells = [c.strip().replace("**", "") for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            value = row[c_idx] if c_idx < len(row) else ""
            para = cell.paragraphs[0]
            run = para.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            if r_idx == 0:
                run.bold = True
    doc.add_paragraph()


def add_markdown(doc: Document, text: str) -> None:
    lines = clean_text(text).splitlines()
    in_code = False
    code_buffer: list[str] = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                if code_buffer:
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_buffer))
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip().startswith("|"):
            rows, next_i = collect_markdown_table(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif re.match(r"^\s*\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\s*\d+\.\s+", "", line), style="List Number")
        elif line.strip().startswith(">"):
            p = doc.add_paragraph(line.strip().lstrip(">").strip())
            p.runs[0].italic = True
        elif line.strip() == "---":
            doc.add_paragraph()
        else:
            cleaned = line.replace("**", "").replace("__", "").replace("`", "")
            doc.add_paragraph(cleaned)
        i += 1


def build_docx(filename: str, title: str, files: list[Path]) -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, title)
    for idx, path in enumerate(files):
        if idx:
            doc.add_page_break()
        doc.add_heading(path.name.replace(".md", "").replace("_", " ").title(), level=1)
        add_markdown(doc, path.read_text(encoding="utf-8"))
    doc.save(OUT / filename)


def build_readme_docx() -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc, "Read Me First - Professor Share Package")
    doc.add_heading("Package Purpose", level=1)
    doc.add_paragraph(
        "This folder contains research-only documents prepared for faculty review. "
        "It includes the problem framing, literature review, dataset study, research gap, "
        "architecture, methodology, evaluation plan, ethics/legal notes, paper outline, "
        "and introduction-writing materials."
    )
    doc.add_heading("What Is Included", level=1)
    items = [
        "Word documents for the main research sections.",
        "Excel sheets for literature, datasets, research gap, objectives, and introduction evidence.",
        "No Markdown files are included in this share folder.",
        "No source-code, run-log, test-output, or implementation artifact is included.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Recommended Reading Order", level=1)
    order = [
        "01_Research_Overview_Problem_Gap.docx",
        "02_Literature_Review_Existing_Models.docx",
        "03_Dataset_Study.docx",
        "04_Architecture_Methodology_Evaluation.docx",
        "05_Ethics_Paper_Outline_Learning_Plan.docx",
        "06_Introduction_Writing_Package.docx",
        "07_Weekly_Update_To_Professor.docx",
    ]
    for item in order:
        doc.add_paragraph(item, style="List Number")
    doc.save(OUT / "00_READ_ME_FIRST.docx")


def autosize(ws) -> None:
    for col in ws.columns:
        width = 12
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value is not None:
                width = max(width, min(len(str(cell.value)) + 2, 65))
        ws.column_dimensions[col_letter].width = width
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")


def write_rows(ws, rows: list[list[str]]) -> None:
    for row in rows:
        ws.append(row)
    autosize(ws)


def read_csv(path: Path) -> list[list[str]]:
    with path.open(newline="", encoding="utf-8") as fh:
        return [[clean_text(cell) for cell in row] for row in csv.reader(fh)]


def build_literature_workbook() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Literature_Matrix"
    write_rows(ws, read_csv(ROOT / "02_literature_matrix.csv"))
    wb.save(OUT / "08_Literature_Matrix.xlsx")


def build_dataset_workbook() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Dataset_Matrix"
    write_rows(ws, read_csv(ROOT / "04_dataset_matrix.csv"))
    shortlist = ROOT / "results/tables/dataset_shortlist.csv"
    if shortlist.exists():
        ws2 = wb.create_sheet("Dataset_Shortlist")
        write_rows(ws2, read_csv(shortlist))
    wb.save(OUT / "09_Dataset_Matrix.xlsx")


def build_research_summary_workbook() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Problem_Understanding"
    write_rows(
        ws,
        [
            ["Section", "Details"],
            [
                "Research topic",
                "SEBA-XAI: Secure Explainable Blockchain-Audited Access Overlay for Inter-Agency Police Data Sharing in India",
            ],
            [
                "Main framing",
                "A research direction for auditable, privacy-aware, explainable access governance over sensitive police-style records.",
            ],
            [
                "Baseline context",
                "CCTNS/ICJS-style infrastructure already exists; the proposed work complements it rather than replacing it.",
            ],
            [
                "Main decision",
                "When an officer or agency requests sensitive data, should the system allow, deny, or escalate the request?",
            ],
            [
                "Scope boundary",
                "No real police data, no deployment claim, no legal-compliance claim, and no individual crime-prediction claim.",
            ],
        ],
    )

    ws = wb.create_sheet("Research_Gap")
    write_rows(
        ws,
        [
            ["Item", "Explanation"],
            [
                "Gap",
                "Existing work often treats blockchain audit, ABAC/PBAC access control, privacy, and XAI separately. The research gap is a joint access-governance framework for sensitive inter-station/inter-agency records.",
            ],
            [
                "Rejected angle 1",
                "Putting raw police records on blockchain is weak because it creates confidentiality, correction, and retention risks.",
            ],
            [
                "Rejected angle 2",
                "Using public NCRB aggregate data for individual suspect prediction is not justified.",
            ],
            [
                "Rejected angle 3",
                "Replacing CCTNS/ICJS is unrealistic and not the research objective.",
            ],
        ],
    )

    ws = wb.create_sheet("Simple_Objectives")
    write_rows(
        ws,
        [
            ["Objective", "Simple explanation"],
            ["O1", "Design an overlay that keeps sensitive records off-chain and logs only audit commitments."],
            ["O2", "Use ABAC/PBAC-style rules to decide allow, deny, or escalate for access requests."],
            ["O3", "Use blockchain-style audit to make access events tamper-evident."],
            ["O4", "Use XAI to explain why an access decision was made."],
            ["O5", "Evaluate security, privacy, latency, auditability, and explanation quality."],
        ],
    )

    ws = wb.create_sheet("RQ_Terms")
    write_rows(
        ws,
        [
            ["Term", "Simple definition"],
            ["Auditability", "Ability to later check who accessed what, why, when, and under which policy."],
            ["Privacy", "Protection of sensitive records and personal information from unnecessary exposure."],
            ["Latency", "Time added by access-control, explanation, logging, and audit steps."],
            ["Metadata exposure", "Risk that non-content information such as time, station, role, or access pattern reveals sensitive facts."],
            ["ABAC", "Access control using subject, object, action, and environment attributes."],
            ["PBAC", "Access control based on explicit policy rules and policy versions."],
            ["XAI", "Explanation methods that make access decisions understandable to officers, approvers, and auditors."],
            ["On-chain commitment", "A hash or proof stored on the ledger instead of raw sensitive data."],
            ["Off-chain record", "Sensitive content stored outside the ledger in protected storage."],
        ],
    )

    ws = wb.create_sheet("Next_Steps_For_Intro")
    write_rows(
        ws,
        [
            ["Step", "Task"],
            ["1", "Finalize the introduction problem statement in India-specific context."],
            ["2", "Use official CCTNS/ICJS sources to establish the baseline."],
            ["3", "Explain the sensitive access-governance problem."],
            ["4", "Justify blockchain, security/access control, and XAI as equal pillars."],
            ["5", "State clearly that the paper does not claim crime prediction or system replacement."],
            ["6", "Write contribution bullets as design and evaluation contributions."],
        ],
    )

    wb.save(OUT / "10_Research_Gap_Objectives_Terms.xlsx")


def build_introduction_evidence_workbook() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Claim_Source_Table"
    write_rows(ws, read_csv(ROOT / "papers/final_paper/introduction/claim_source_table.csv"))
    ws2 = wb.create_sheet("Evidence_Register")
    write_rows(ws2, read_csv(ROOT / "papers/final_paper/introduction/evidence_register.csv"))
    wb.save(OUT / "11_Introduction_Evidence.xlsx")


def copy_focused_workbook() -> None:
    src = ROOT / "spreadsheets/SEBA_XAI_Focused_Research_Only.xlsx"
    if src.exists():
        dst = OUT / "12_Focused_Research_Only.xlsx"
        shutil.copy2(src, dst)
        wb = load_workbook(dst)
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        cell.value = clean_text(cell.value)
            autosize(ws)
        wb.save(dst)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_readme_docx()
    for filename, title, files in DOC_SETS:
        build_docx(filename, title, files)
    build_literature_workbook()
    build_dataset_workbook()
    build_research_summary_workbook()
    build_introduction_evidence_workbook()
    copy_focused_workbook()
    print(f"wrote share package: {OUT}")


if __name__ == "__main__":
    main()
